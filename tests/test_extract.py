"""文本抽取测试"""
from pathlib import Path
import docx
from outline_extraction.parsing.extract import extract_document, _is_scanned


def test_extract_docx(tmp_path):
    """docx 抽取为 Markdown，方法标记 docx"""
    p = tmp_path / "a.docx"
    d = docx.Document()
    d.add_heading("第一章 总则", level=1)
    d.add_paragraph("这是正文。")
    d.save(p)
    result = extract_document(Path(p), ".docx")
    assert result.extract_method == "docx"
    assert "第一章 总则" in result.raw_markdown
    assert "这是正文。" in result.raw_markdown


def test_extract_docx_heading_becomes_markdown(tmp_path):
    """docx 标题转成 Markdown # 前缀"""
    p = tmp_path / "h.docx"
    d = docx.Document()
    d.add_heading("标题A", level=2)
    d.save(p)
    result = extract_document(Path(p), ".docx")
    assert "## 标题A" in result.raw_markdown


def test_is_scanned_threshold():
    """每页平均字符数低于阈值判为扫描件"""
    assert _is_scanned(total_chars=50, page_count=10) is True   # 5 字/页
    assert _is_scanned(total_chars=5000, page_count=10) is False  # 500 字/页


def test_extract_unknown_suffix(tmp_path):
    """未知后缀返回空 markdown 并标记 skipped，不报错"""
    p = tmp_path / "x.bin"
    p.write_bytes(b"\x00\x01")
    result = extract_document(Path(p), ".bin")
    assert result.extract_method == "skipped"
    assert result.raw_markdown == ""


def test_docx_table_stays_in_place(tmp_path):
    """表格应保留在其所属段落之间，而非被甩到文末"""
    p = tmp_path / "t.docx"
    d = docx.Document()
    d.add_heading("评分办法", level=1)
    d.add_paragraph("评分标准如下：")
    tbl = d.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "评分项"; tbl.cell(0, 1).text = "分值"
    tbl.cell(1, 0).text = "ISO9001"; tbl.cell(1, 1).text = "2分"
    d.add_paragraph("以上为评分细则。")
    d.save(p)
    md = extract_document(Path(p), ".docx").raw_markdown
    lines = md.splitlines()
    idx_intro = next(i for i, l in enumerate(lines) if "评分标准如下" in l)
    idx_table = next(i for i, l in enumerate(lines) if "ISO9001" in l)
    idx_after = next(i for i, l in enumerate(lines) if "以上为评分细则" in l)
    assert idx_intro < idx_table < idx_after
    assert any(set(l.strip()) <= set("| -") and "-" in l for l in lines)


def test_docx_table_has_header_separator(tmp_path):
    """表格首行后应有 | --- | 分隔行"""
    p = tmp_path / "t2.docx"
    d = docx.Document()
    tbl = d.add_table(rows=2, cols=3)
    for c in range(3):
        tbl.cell(0, c).text = f"列{c}"
        tbl.cell(1, c).text = f"值{c}"
    d.save(p)
    md = extract_document(Path(p), ".docx").raw_markdown
    lines = [l for l in md.splitlines() if l.startswith("|")]
    assert len(lines) >= 3
    assert "---" in lines[1]
