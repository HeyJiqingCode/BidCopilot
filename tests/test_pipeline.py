"""端到端管线测试——真实解析 + fake LLM"""
from pathlib import Path
import docx
from bid_copilot.models import (
    OutlineNode, SourceRef, SourceType, RequirementItem,
)
from bid_copilot.understanding.classify import FileClass, ClassifyResult
from bid_copilot.understanding.locate import LocateResult
from bid_copilot.understanding.extract_skeleton import SkeletonResult
from bid_copilot.understanding.extract_requirements import RequirementsResult
from bid_copilot.understanding.alignment.merge import (
    MergeResult, MergeDecision, Disposition, _NormalizeResult, _AttachResult, _DedupeResult,
)
from bid_copilot.understanding.alignment.supplement import SupplementResult
from bid_copilot.understanding.pipeline import run_pipeline


def _nodes_from_skeleton(kwargs):
    """从阶段A 的 input_content 还原骨架 OutlineNode 列表——测试辅助（规范化原样返回骨架）"""
    import json
    from bid_copilot.models import OutlineNode
    try:
        payload = json.loads(kwargs.get("input_content", "[]"))
    except Exception:
        payload = []
    return [OutlineNode(**n) for n in payload] if isinstance(payload, list) else []


class _ScriptedLLM:
    """脚本化 LLM：merge 三阶段（normalize/attach/dedupe）按 schema 自动应答，其余按 push 顺序返回。

    pipeline 测试只 push 业务步骤（classify/locate/skeleton/requirements/supplement）的返回，
    merge 内部拆成的多次调用由本类按 schema 兜底，测试无需关心 merge 调了几次。
    """
    def __init__(self):
        self.script = []
        self.idx = 0

    def push(self, result):
        self.script.append(result)

    def complete(self, **kwargs):
        schema = kwargs.get("schema")
        if schema is _NormalizeResult:                       # 阶段A：原样返回骨架
            return _NormalizeResult(tree=_nodes_from_skeleton(kwargs))
        if schema is _AttachResult:                          # 阶段B：空判定
            return _AttachResult()
        if schema is _DedupeResult:                          # 阶段C：空分组
            return _DedupeResult()
        r = self.script[self.idx]
        self.idx += 1
        return r


def _make_docx(path):
    d = docx.Document()
    d.add_heading("投标文件格式", level=1)
    d.add_heading("投标函", level=2)
    d.save(path)


def test_run_pipeline_end_to_end(tmp_path):
    """管线跑通：解析真实 docx + 脚本化 LLM，产出 OutlineTree"""
    src = tmp_path / "fmt.docx"
    _make_docx(src)

    llm = _ScriptedLLM()
    # classify（1个文档）
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    # locate
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[],
                          tech_spec_sections=[], business_sections=[]))
    # extract_skeleton
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[
        SourceRef(type=SourceType.SKELETON, document="fmt.docx", location="一", quote=None)], children=[])]))
    # merge 内部三阶段由 _ScriptedLLM 按 schema 兜底（规范化原样返回骨架、无要求可挂），无需 push
    # supplement：无游离要求会被跳过，这里 push 空判定占位（不会被消费）
    llm.push(SupplementResult())

    steps_seen = []
    tree = run_pipeline(src, llm=llm, model_main="gpt-5.4", model_mini="gpt-5.4-mini",
                        run_dir=tmp_path / "run", progress_callback=lambda s, p: steps_seen.append(s))

    assert tree.nodes[0].title == "投标函"
    assert tree.nodes[0].id == "1"           # 已 finalize id
    assert "parse" in steps_seen
    assert "merge" in steps_seen


def test_run_pipeline_accepts_explicit_project_name(tmp_path):
    """run_pipeline 可显式指定 project_name，覆盖默认的 input_path.stem"""
    src = tmp_path / "input"
    src.mkdir()
    f = src / "fmt.docx"
    import docx
    d = docx.Document()
    d.add_heading("投标文件格式", level=1)
    d.save(f)

    llm = _ScriptedLLM()
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[], tech_spec_sections=[], business_sections=[]))
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[], children=[])]))
    # merge 三阶段由 _ScriptedLLM 按 schema 兜底
    llm.push(SupplementResult())   # supplement 现只返回安置判定；这些用例无游离要求，不会被消费

    tree = run_pipeline(src, llm=llm, model_main="m", model_mini="mini",
                        run_dir=tmp_path / "run", project_name="淮能项目")
    assert tree.project_name == "淮能项目"


def test_pipeline_emits_detail_level_logs(tmp_path):
    """管线应发出带 level=detail 的细粒度日志事件（不止 main 摘要）"""
    src = tmp_path / "input"
    src.mkdir()
    f = src / "fmt.docx"
    import docx
    d = docx.Document()
    d.add_heading("投标文件格式", level=1)
    d.save(f)

    llm = _ScriptedLLM()
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[], tech_spec_sections=[], business_sections=[]))
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[], children=[])]))
    # merge 三阶段由 _ScriptedLLM 按 schema 兜底
    llm.push(SupplementResult())   # supplement 现只返回安置判定；这些用例无游离要求，不会被消费

    events = []
    run_pipeline(src, llm=llm, model_main="gpt-5.4", model_mini="gpt-5.4-mini",
                 run_dir=tmp_path / "run", log_callback=lambda e: events.append(e))

    # 每个事件都带 level 字段
    assert all("level" in e for e in events)
    # 至少有一条 detail 级日志（细粒度），且提到模型或文件
    details = [e for e in events if e.get("level") == "detail"]
    assert len(details) >= 1
    # classify 阶段应有提到分类的 detail 日志
    assert any(e["phase"] == "classify" and e["level"] == "detail" for e in events)


def test_pipeline_detail_logs_cover_main_phases(tmp_path):
    """parse/locate/extract_skeleton/extract_requirements/merge 都应有 detail 日志"""
    src = tmp_path / "input"
    src.mkdir()
    f = src / "fmt.docx"
    import docx
    d = docx.Document()
    d.add_heading("投标文件格式", level=1)
    d.add_paragraph("一、商务投标文件")
    d.save(f)

    from bid_copilot.understanding.extract_requirements import RequirementsResult
    from bid_copilot.models import RequirementItem

    llm = _ScriptedLLM()
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[0], tech_spec_sections=[], business_sections=[]))
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[], children=[])]))
    llm.push(RequirementsResult(items=[RequirementItem(ref_id="", description="x", source_type=SourceType.SCORING, location="评1", suggested_title="X")]))
    # merge 三阶段由 _ScriptedLLM 按 schema 兜底
    llm.push(SupplementResult())   # supplement 现只返回安置判定；这些用例无游离要求，不会被消费

    events = []
    run_pipeline(src, llm=llm, model_main="gpt-5.4", model_mini="gpt-5.4-mini",
                 run_dir=tmp_path / "run", log_callback=lambda e: events.append(e))

    phases_with_detail = {e["phase"] for e in events if e.get("level") == "detail"}
    for ph in ["parse", "locate", "extract_skeleton", "extract_requirements", "merge"]:
        assert ph in phases_with_detail, f"{ph} 缺少 detail 日志"


class _RecordingLLM(_ScriptedLLM):
    """记录每次调用所用 model，用于断言档位路由——测试辅助"""
    def __init__(self, sink):
        super().__init__()
        self._sink = sink

    def complete(self, **kwargs):
        self._sink.append(kwargs.get("model"))
        return super().complete(**kwargs)


def _push_minimal(llm):
    """给 scripted LLM 压入最小可跑通脚本（classify/locate/skeleton/supplement）——测试辅助"""
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[], tech_spec_sections=[], business_sections=[]))
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[], children=[])]))
    llm.push(SupplementResult())   # supplement 现只返回安置判定；这些用例无游离要求，不会被消费


def _make_fmt_docx(tmp_path):
    """造一个仅含投标文件格式标题的最小 docx 输入目录——测试辅助"""
    src = tmp_path / "input"
    src.mkdir()
    d = docx.Document()
    d.add_heading("投标文件格式", level=1)
    d.save(src / "fmt.docx")
    return src


def test_pipeline_unplaced_req_goes_to_supplement_even_if_not_floating(tmp_path):
    """两头落空防御：merge 判要求可挂(child_of)但 node_id 无效→未进树时，该要求仍须交 supplement 兜底

    复现原 run R17/R18 的丢失路径：disposition 非 floating，但因 node_id 无效未真正回填进树。
    旧逻辑按 disposition 挑 floating 会漏掉它；修复后按"是否真在树里"判定，必须把它交给 supplement。
    """
    src = _make_fmt_docx(tmp_path)
    seen_floating = []

    from bid_copilot.understanding.alignment.supplement import (
        SupplementDecision, SupplementDisposition,
    )

    class _InvalidNodeIdLLM(_ScriptedLLM):
        """attach 阶段对要求给出 child_of 但 node_id 无效；supplement 阶段把它正确安置到 "1"。"""
        def complete(self, **kwargs):
            schema = kwargs.get("schema")
            if schema is _AttachResult:
                # 判 R0 可挂到一个不存在的节点 → 工程回填会静默跳过 → R0 不进树
                return _AttachResult(decisions=[
                    MergeDecision(ref_id="R0", disposition=Disposition.CHILD_OF, node_id="NOPE_INVALID")])
            if schema is SupplementResult:
                import json
                payload = json.loads(kwargs.get("input_content", "{}"))
                # 新 schema：floating_requirements 是对象列表，记录其 description
                seen_floating.extend(f["description"] for f in payload.get("floating_requirements", []))
                # 把 R0 安置到已有节点 "1"（ref_id 由工程回填）
                return SupplementResult(decisions=[
                    SupplementDecision(ref_id="R0", disposition=SupplementDisposition.MERGED_INTO, node_id="1")])
            return super().complete(**kwargs)

    llm = _InvalidNodeIdLLM()
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[0], tech_spec_sections=[], business_sections=[]))
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[], children=[])]))
    llm.push(RequirementsResult(items=[RequirementItem(
        ref_id="", description="履约保证金退还承诺", source_type=SourceType.BIZ_TERMS,
        location="22.3", suggested_title="履约保证金")]))

    tree = run_pipeline(src, llm=llm, model_main="gpt-5.4", model_mini="gpt-5.4-mini", run_dir=tmp_path / "run")

    # 该要求未挂进树（node_id 无效），但被识别为未安置 → 交给了 supplement
    assert "履约保证金退还承诺" in seen_floating
    # 且 supplement 工程回填后，覆盖率不再漏它（端到端验证 A 修复）
    assert tree.coverage.mapped_biz_items == tree.coverage.total_biz_items
    assert "履约保证金退还承诺" not in tree.coverage.unmapped


def test_pipeline_skips_supplement_when_no_floating(tmp_path):
    """无游离要求时跳过生成式补充——不发起 SupplementResult 的 LLM 调用"""
    src = _make_fmt_docx(tmp_path)
    schemas_seen = []

    class _SchemaRecordingLLM(_ScriptedLLM):
        """记录每次调用的 schema，用于断言 supplement 是否被调用"""
        def complete(self, **kwargs):
            schemas_seen.append(kwargs.get("schema"))
            return super().complete(**kwargs)

    llm = _SchemaRecordingLLM()
    # 不 push 任何要求 → merge 无要求可挂 → 0 条游离 → supplement 应被跳过
    llm.push(ClassifyResult(file_class=FileClass.BID_FORMAT, confidence=0.9))
    llm.push(LocateResult(bid_format_sections=[0], scoring_sections=[], tech_spec_sections=[], business_sections=[]))
    llm.push(SkeletonResult(nodes=[OutlineNode(id="1", title="投标函", level=1, sources=[], children=[])]))

    tree = run_pipeline(src, llm=llm, model_main="MAIN", model_mini="MINI", run_dir=tmp_path / "run")

    assert SupplementResult not in schemas_seen   # 0 游离 → 未调用生成式补充
    assert tree.nodes[0].title == "投标函"          # 树仍正常产出（沿用归并树）


def test_pipeline_model_tier_routing(tmp_path):
    """models 档位映射：classify 走 mini、merge 显式配 mini 时也走 mini、显式 nano 走 NANO，其余 main"""
    src = _make_fmt_docx(tmp_path)
    used_models = []
    llm = _RecordingLLM(used_models)
    _push_minimal(llm)

    run_pipeline(src, llm=llm, model_main="MAIN", model_mini="MINI", model_nano="NANO",
                 run_dir=tmp_path / "run",
                 models={"classify": "mini", "locate": "main", "merge": "mini", "skeleton": "nano"})

    # classify 用 MINI；locate 用 MAIN；merge 显式配 mini → MINI；skeleton 显式配 nano → NANO
    assert "MINI" in used_models    # classify / merge
    assert "MAIN" in used_models    # locate 等
    assert "NANO" in used_models    # skeleton 显式 nano 档


def test_pipeline_nano_falls_back_to_mini_when_unset(tmp_path):
    """nano 档但未配置 model_nano（空串）时回退到 model_mini，不传空模型名"""
    src = _make_fmt_docx(tmp_path)
    used_models = []
    llm = _RecordingLLM(used_models)
    _push_minimal(llm)

    # skeleton 配 nano，但 model_nano 缺省为空串 → 应回退用 MINI
    run_pipeline(src, llm=llm, model_main="MAIN", model_mini="MINI",
                 run_dir=tmp_path / "run",
                 models={"skeleton": "nano"})

    assert "" not in used_models    # 不应出现空模型名
    assert "MINI" in used_models    # nano 回退到 mini
