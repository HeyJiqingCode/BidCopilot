"""文本抽取——探测优先决策链，统一输出 Markdown"""
import subprocess
from pathlib import Path
from typing import Any
import docx
import pdfplumber
from outline_extraction.models import ParsedDocument
from outline_extraction.parsing.cu_client import analyze_with_cu

# 扫描件判定阈值：每页平均字符数低于此值视为扫描件
_SCANNED_CHARS_PER_PAGE = 50


def extract_document(file_path: Path, suffix: str, cu: Any = None) -> ParsedDocument:
    """按后缀分流抽取文本

    分流规则：.docx 走本地抽取（保留 Word 原生标题层级）；其余后缀优先用 CU 出结构化
    markdown，CU 不可用或失败时降级回本地（textutil/pdfplumber/xml）。

    参数:
        file_path: 文件路径
        suffix: 小写后缀
        cu: CU 客户端（None 表示未配置，走本地兜底）
    返回:
        ParsedDocument（含统一 Markdown 与抽取方式标记）
    """
    name = file_path.name
    if suffix == ".docx":
        return ParsedDocument(filename=name, raw_markdown=_docx_to_md(file_path),
                              extract_method="docx", page_count=None)
    if cu is not None:
        try:
            cu_result = analyze_with_cu(file_path, cu)
            if cu_result.markdown.strip():
                return ParsedDocument(filename=name, raw_markdown=cu_result.markdown,
                                      extract_method="cu", page_count=cu_result.page_count)
        except Exception:
            pass
    if suffix == ".doc":
        return ParsedDocument(filename=name, raw_markdown=_doc_to_md(file_path),
                              extract_method="textutil", page_count=None)
    if suffix == ".pdf":
        return _pdf_to_doc(file_path, name)
    if suffix == ".xml":
        return ParsedDocument(filename=name, raw_markdown=file_path.read_text(errors="ignore"),
                              extract_method="xml", page_count=None)
    return ParsedDocument(filename=name, raw_markdown="", extract_method="skipped", page_count=None)


def _docx_to_md(file_path: Path) -> str:
    """docx → Markdown：按文档真实顺序遍历段落与表格——内部辅助

    段落 Heading 样式转 # 前缀；表格在其原始位置转 Markdown 表格（含表头分隔行），
    不再统一追加到文末，从而保住表格与所属章节的相邻关系（评分表/技术参数表常为表格）。

    参数:
        file_path: docx 文件路径
    返回:
        统一 Markdown 文本
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(str(file_path))
    lines: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "") if para.style else ""
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                lines.append("#" * min(level, 6) + " " + text)
            else:
                lines.append(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            lines.extend(_table_to_md(table))
    return "\n".join(lines)


def _table_to_md(table) -> list[str]:
    """把一个 docx 表格转成 Markdown 表格行（含表头分隔行）——内部辅助

    参数:
        table: python-docx Table 对象
    返回:
        Markdown 表格行列表；空表返回空列表
    """
    rows = table.rows
    if not rows:
        return []
    md_rows: list[str] = []
    for r_idx, row in enumerate(rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if r_idx == 0:
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return md_rows


def _doc_to_md(file_path: Path) -> str:
    """.doc → 文本：调用 macOS textutil 子进程——内部辅助"""
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(file_path)],
        capture_output=True, text=True,
    )
    return result.stdout


def _pdf_to_doc(file_path: Path, name: str, cu: Any = None) -> ParsedDocument:
    """PDF：抽文本层；字符过少视为扫描件，优先走 CU，CU 不可用则降级——内部辅助"""
    text_parts: list[str] = []
    page_count = 0
    with pdfplumber.open(str(file_path)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    full = "\n".join(text_parts)
    if _is_scanned(len(full.strip()), page_count):
        cu_result = analyze_with_cu(file_path, cu)
        if cu_result.markdown:                       # CU 可用：用其结构化输出
            return ParsedDocument(filename=name, raw_markdown=cu_result.markdown,
                                  extract_method="cu_ocr",
                                  page_count=cu_result.page_count or page_count)
        return ParsedDocument(filename=name, raw_markdown=full,    # 降级：保留原文本
                              extract_method="needs_ocr", page_count=page_count)
    return ParsedDocument(filename=name, raw_markdown=full,
                          extract_method="pdf_text", page_count=page_count)


def _is_scanned(total_chars: int, page_count: int) -> bool:
    """判断 PDF 是否为扫描件——与具体文件无关的通用指标

    参数:
        total_chars: 抽到的总字符数
        page_count: 页数
    返回:
        True 表示疑似扫描件（每页平均字符数过低）
    """
    if page_count <= 0:
        return total_chars == 0
    return (total_chars / page_count) < _SCANNED_CHARS_PER_PAGE
