"""Word 导出——OutlineTree 遍历为多级标题 docx"""
from pathlib import Path
import docx
from bid_copilot.models import OutlineTree, OutlineNode


def export_to_docx(tree: OutlineTree, output_path: Path) -> Path:
    """把大纲树导出为 Word 多级标题文档

    参数:
        tree: 完整大纲树
        output_path: 输出 .docx 路径
    返回:
        输出路径
    """
    doc = docx.Document()
    doc.add_heading(tree.project_name, level=0)
    for node in tree.nodes:
        _write_node(doc, node)
    doc.save(str(output_path))
    return output_path


def _write_node(doc, node: OutlineNode) -> None:
    """递归写入单个节点为 Heading——内部辅助

    参数:
        doc: python-docx Document
        node: 当前节点
    返回: 无
    """
    # 标题前带路径式序号（如 "1.2 法定代表人身份证明"），与网页大纲一致
    title = f"{node.id} {node.title}"
    doc.add_heading(title, level=min(node.level, 9))
    for child in node.children:
        _write_node(doc, child)
